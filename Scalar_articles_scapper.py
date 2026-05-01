from docx import Document
import requests
from bs4 import BeautifulSoup
url = ['https://www.scaler.com/topics/second-largest-number-in-python/']
def save_articles_to_word(urls):
    doc = Document()
    doc.add_heading('Python Tutorial Articles', 0)

    for url in urls:
        # 1. Fetch the page
        res = requests.get(url)
        soup = BeautifulSoup(res.text, 'html.parser')
        
        # 2. Extract Title and Text (Selectors depend on the site's HTML)
        title = soup.find('h1').text
        content = soup.find('div', {'class': 'article-content'}).text 

        # 3. Add to Word Document
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)
        doc.add_page_break()

    doc.save('Python_Tutorial.docx')